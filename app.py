import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import matplotlib.colors as mcolors
import math
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from shapely.geometry import Point, LineString, Polygon
import shapely.affinity as affinity
from matplotlib.path import Path
from streamlit_drawable_canvas import st_canvas

# ==========================================
# CẤU HÌNH TRANG WEB
# ==========================================
st.set_page_config(page_title="Riken Viet - Enterprise Gas Mapping", layout="wide")
st.title("🛡️ Riken Viet - Hệ thống Thiết kế Vùng phủ Khí Đa lớp (3D)")

if 'room_data' not in st.session_state:
    st.session_state.room_data = pd.DataFrame({"X": [0, 15, 15, 0, 0], "Y": [0, 0, 10, 10, 0]}) 
if 'obs_data' not in st.session_state:
    st.session_state.obs_data = pd.DataFrame([
        {"Type": "Cylinder", "X": 7.5, "Y": 5.0, "Width_Radius": 1.5, "Length": 0.0, "Height": 4.0, "Angle": 0}
    ])
if 'auto_config' not in st.session_state:
    st.session_state.auto_config = pd.DataFrame([
        {"Target Gas": "CH4", "Layer": "Khí Nhẹ (Sát trần)", "Model": "SD-1 (Catalytic)", "Radius": 5.0, "Color": "cyan"},
        {"Target Gas": "H2S", "Layer": "Khí Nặng (Sát sàn)", "Model": "GD-70D (Electro)", "Radius": 4.0, "Color": "magenta"}
    ])
if 'det_data' not in st.session_state:
    st.session_state.det_data = pd.DataFrame(columns=["ID", "Gas", "X", "Y", "Z", "Radius", "Color"])

# ==========================================
# 1. GIAO DIỆN NHẬP LIỆU & BẢNG VẼ TƯƠNG TÁC
# ==========================================
col_input1, col_input2 = st.columns([1.2, 1])

with col_input1:
    st.header("1. Kiến trúc Phòng Không gian")
    room_z = st.number_input("Chiều cao trần (Z) - mét", min_value=1.0, value=5.0)

    st.subheader("🖍️ Bảng vẽ Mặt bằng (Auto-Snap)")
    st.markdown("""
    **Mẹo vẽ chuẩn xác:** 1. Chọn công cụ vẽ Đa giác (Polygon) ở thanh công cụ bên trái.
    2. Click để chấm các góc tường. Đừng lo vẽ lệch, AI sẽ tự nắn thẳng.
    3. 🚨 **LƯU Ý:** Để khép kín căn phòng, hãy **Click ĐÚP CHUỘT (Double-click)** ở điểm cuối cùng!
    """)
    
    canvas_w, canvas_h, grid_size = 600, 400, 20 # Quy ước: 20 pixels = 1 mét
    
    canvas_result = st_canvas(
        fill_color="rgba(0, 255, 150, 0.3)",  
        stroke_width=3,
        stroke_color="#00FFAA",
        background_color="#1E1E1E", 
        height=canvas_h, width=canvas_w,
        drawing_mode="polygon",
        key="canvas",
    )

    # NÚT ĐỒNG BỘ NÉT VẼ
    if st.button("📥 Bấm vào đây để Đồng bộ nét vẽ xuống Tọa độ", type="primary", use_container_width=True):
        if canvas_result.json_data is not None:
            objects = canvas_result.json_data["objects"]
            if len(objects) > 0:
                last_obj = objects[-1]
                if "points" in last_obj:
                    pts = last_obj["points"]
                    drawn_room = []
                    for p in pts:
                        # Tọa độ tuyệt đối: Lấy vị trí click chia cho tỷ lệ mét
                        real_x = p["x"] / grid_size
                        real_y = p["y"] / grid_size
                        
                        # AUTO-SNAP: Nắn về bội số của 0.5m
                        snap_x = round(real_x * 2) / 2
                        snap_y = round(real_y * 2) / 2
                        drawn_room.append({"X": snap_x, "Y": snap_y})
                    
                    # Tự động khép vòng tọa độ nếu điểm đầu khác điểm cuối
                    if drawn_room[0] != drawn_room[-1]:
                        drawn_room.append(drawn_room[0])
                    
                    df_drawn = pd.DataFrame(drawn_room)
                    # Lọc các điểm click đúp trùng nhau
                    df_drawn = df_drawn.loc[(df_drawn.shift() != df_drawn).any(axis=1)].reset_index(drop=True)
                    st.session_state.room_data = df_drawn
                    st.success(f"✅ Đã nhận diện thành công phòng đa giác với {len(df_drawn)-1} góc!")
                else:
                    st.error("Chưa nhận diện được hình khối. Nhớ 'Click đúp chuột' khi vẽ xong nhé!")
            else:
                st.warning("Bảng vẽ đang trống. Hãy vẽ thử một hình trước khi đồng bộ!")

    col_r1, col_r2 = st.columns([1, 1])
    with col_r1:
        if st.button("🔄 Tạo nhanh phòng Chữ Nhật", use_container_width=True):
            st.session_state.room_data = pd.DataFrame({"X": [0, 15, 15, 0, 0], "Y": [0, 0, 10, 10, 0]})
            st.rerun()
    with col_r2:
        if st.button("🔄 Tạo mẫu phòng Chữ L", use_container_width=True):
            # Tạo mẫu phòng chữ L điển hình (15m x 15m khuyết góc)
            st.session_state.room_data = pd.DataFrame({"X": [0, 15, 15, 7, 7, 0, 0], "Y": [0, 0, 8, 8, 15, 15, 0]})
            st.rerun()

    st.caption("👇 Bảng tọa độ không gian thực (có thể chỉnh sửa thủ công nếu AI bắt điểm hơi lệch):")
    edited_room = st.data_editor(st.session_state.room_data, num_rows="dynamic", use_container_width=True)
    
    if len(edited_room) >= 3:
        room_poly = Polygon(list(zip(edited_room['X'], edited_room['Y'])))
    else:
        st.error("Phòng cần ít nhất 3 góc (tọa độ)!")
        room_poly = None

with col_input2:
    st.header("2. Cấu hình Thiết bị & Bố trí")
    
    with st.expander("🚧 Danh sách Vật cản (Cylinder / Box)", expanded=True):
        edited_obs = st.data_editor(
            st.session_state.obs_data, num_rows="dynamic", use_container_width=True,
            column_config={"Type": st.column_config.SelectboxColumn("Loại", options=["Cylinder", "Box"])}
        )
        st.session_state.obs_data = edited_obs

    with st.expander("⚙️ Thiết lập Các Phân hệ Khí (Bấm '+' để thêm)", expanded=True):
        edited_auto_config = st.data_editor(
            st.session_state.auto_config, num_rows="dynamic", use_container_width=True,
            column_config={
                "Layer": st.column_config.SelectboxColumn("Mặt phẳng", options=["Khí Nhẹ (Sát trần)", "Khí Trung bình (Vùng thở)", "Khí Nặng (Sát sàn)"]),
                "Color": st.column_config.SelectboxColumn("Màu bản đồ", options=["cyan", "magenta", "yellow", "lime", "red", "blue", "orange"])
            }
        )
        st.session_state.auto_config = edited_auto_config

        if st.button("🚀 Tự động Rải Đầu dò theo Bảng cấu hình trên", type="primary"):
            if edited_auto_config.empty:
                st.warning("Vui lòng thiết lập ít nhất 1 loại khí!")
            elif room_poly is not None:
                new_dets = []
                for _, row_cfg in edited_auto_config.iterrows():
                    if "Nhẹ" in row_cfg["Layer"]: z_val = max(room_z - 0.5, 0.5)
                    elif "Nặng" in row_cfg["Layer"]: z_val = 0.5
                    else: z_val = 1.5
                    
                    spacing = row_cfg["Radius"] * 1.5 
                    minx, miny, maxx, maxy = room_poly.bounds
                    nx = max(1, math.ceil((maxx - minx) / spacing))
                    ny = max(1, math.ceil((maxy - miny) / spacing))
                    
                    x_steps = np.linspace(minx + (maxx-minx)/(2*nx), maxx - (maxx-minx)/(2*nx), nx)
                    y_steps = np.linspace(miny + (maxy-miny)/(2*ny), maxy - (maxy-miny)/(2*ny), ny)
                    
                    count = 1
                    for x in x_steps:
                        for y in y_steps:
                            pt = Point(x, y)
                            # ĐIỂM CHỐT AI: Chỉ rải nếu điểm x, y thực sự nằm "Bên Trong" phòng đa giác
                            if room_poly.contains(pt): 
                                new_dets.append({
                                    "ID": f"{row_cfg['Model']} ({count:02d})", 
                                    "Gas": f"{row_cfg['Target Gas']} ({row_cfg['Layer'].split(' ')[1]})", 
                                    "X": round(x, 1), "Y": round(y, 1),
                                    "Z": z_val, "Radius": row_cfg["Radius"], "Color": row_cfg["Color"]
                                })
                                count += 1
                
                st.session_state.det_data = pd.DataFrame(new_dets)
                st.success(f"Đã rải thành công {len(new_dets)} đầu dò. Lưới rải đã tự động tránh vùng khuyết của đa giác!")
                st.rerun()

    st.write("📋 **Bảng Tọa độ Đầu dò Thực tế (Chỉnh sửa thủ công nếu bị vướng):**")
    edited_dets = st.data_editor(
        st.session_state.det_data, num_rows="dynamic", use_container_width=True,
        column_config={
            "Color": st.column_config.SelectboxColumn("Màu", options=["cyan", "magenta", "yellow", "lime", "red", "white", "blue", "orange"])
        }
    )
    st.session_state.det_data = edited_dets


# ==========================================
# 2. CÁC HÀM LÕI TOÁN HỌC & ĐỒ HỌA
# ==========================================
def create_obstacle_polys(df_obs):
    obs_polys = []
    for _, row in df_obs.iterrows():
        if row['Type'] == 'Cylinder':
            obs_polys.append(Point(row['X'], row['Y']).buffer(row['Width_Radius']))
        elif row['Type'] == 'Box':
            w, l = row['Width_Radius'], row['Length']
            box = Polygon([(row['X']-w/2, row['Y']-l/2), (row['X']+w/2, row['Y']-l/2),
                           (row['X']+w/2, row['Y']+l/2), (row['X']-w/2, row['Y']+l/2)])
            box = affinity.rotate(box, row.get('Angle', 0), origin='center')
            obs_polys.append(box)
    return obs_polys

def check_collision_shapely(df_dets, obs_polys):
    collisions = []
    if not obs_polys or df_dets.empty: return collisions
    for _, det in df_dets.iterrows():
        pt = Point(det['X'], det['Y'])
        if any(poly.contains(pt) for poly in obs_polys):
            collisions.append(det['ID'])
    return collisions

def generate_2d_plot(room_poly, obs_polys, df_dets_group, gas_name):
    minx, miny, maxx, maxy = room_poly.bounds
    res = 0.2
    xx, yy = np.meshgrid(np.arange(minx, maxx, res), np.arange(miny, maxy, res))
    pts_x, pts_y = xx.flatten(), yy.flatten()

    room_path = Path(list(room_poly.exterior.coords))
    in_room_mask = room_path.contains_points(np.c_[pts_x, pts_y])
    
    in_obs_mask = np.zeros(len(pts_x), dtype=bool)
    for obs in obs_polys:
        obs_path = Path(list(obs.exterior.coords))
        in_obs_mask |= obs_path.contains_points(np.c_[pts_x, pts_y])
    
    valid_points_mask = in_room_mask & ~in_obs_mask
    coverage_mask = np.zeros(len(pts_x), dtype=bool)

    for _, det in df_dets_group.iterrows():
        dx, dy, dr = det['X'], det['Y'], det['Radius']
        dist_sq = (pts_x - dx)**2 + (pts_y - dy)**2
        in_radius_mask = dist_sq <= dr**2
        
        check_mask = valid_points_mask & in_radius_mask
        det_pt = Point(dx, dy)
        
        for i in np.where(check_mask)[0]:
            if coverage_mask[i]: continue
            line = LineString([det_pt, Point(pts_x[i], pts_y[i])])
            shadowed = any(line.crosses(obs) for obs in obs_polys)
            if not shadowed: coverage_mask[i] = True

    diem_kha_dung = np.sum(valid_points_mask)
    ty_le = (np.sum(coverage_mask) / diem_kha_dung) * 100 if diem_kha_dung > 0 else 0

    fig, ax = plt.subplots(figsize=(8, 6))
    coverage_2d = coverage_mask.reshape(xx.shape)
    ax.contourf(xx, yy, coverage_2d, levels=[0.5, 1], colors=['#A8E6CF'], alpha=0.6, zorder=1)
    
    rx, ry = room_poly.exterior.xy
    ax.plot(rx, ry, 'k-', lw=3, zorder=2)
    
    for obs in obs_polys:
        ox, oy = obs.exterior.xy
        ax.fill(ox, oy, color='gray', alpha=0.8, zorder=4)

    valid_colors = mcolors.CSS4_COLORS
    for _, det in df_dets_group.iterrows():
        c = det['Color'].lower()
        use_c = c if c in valid_colors else 'blue'
        ax.add_patch(plt.Circle((det['X'], det['Y']), det['Radius'], color=use_c, fill=False, linestyle='--', lw=1.5, zorder=3))
        ax.text(det['X']+0.3, det['Y']+0.3, f"{det['ID']}", fontsize=8, color='black', zorder=6, bbox=dict(facecolor='white', alpha=0.7, edgecolor='none', pad=1))
        ax.plot(det['X'], det['Y'], '^', color=use_c, markersize=12, markeredgecolor='black', zorder=5)

    ax.set_title(f"Bản đồ phân tích: {gas_name} | Mức an toàn: {ty_le:.1f}%", fontweight='bold')
    ax.axis('equal'); ax.grid(True, linestyle=':', alpha=0.5)
    return fig, ty_le

def generate_plotly_3d_complex(room_poly, rz, obs_polys, df_obs, df_dets):
    fig = go.Figure()
    rx, ry = room_poly.exterior.xy
    rx, ry = list(rx), list(ry)
    
    fig.add_trace(go.Scatter3d(x=rx, y=ry, z=[0]*len(rx), mode='lines', line=dict(color='white', width=4), name='Đáy tường'))
    fig.add_trace(go.Scatter3d(x=rx, y=ry, z=[rz]*len(rx), mode='lines', line=dict(color='white', width=4), name='Đỉnh tường'))
    for x, y in zip(rx[:-1], ry[:-1]):
        fig.add_trace(go.Scatter3d(x=[x,x], y=[y,y], z=[0,rz], mode='lines', line=dict(color='white', width=2), showlegend=False))

    def get_sphere(x0, y0, z0, r):
        u, v = np.mgrid[0:2*np.pi:15j, 0:np.pi:10j]
        return r*np.cos(u)*np.sin(v)+x0, r*np.sin(u)*np.sin(v)+y0, r*np.cos(v)+z0

    for _, det in df_dets.iterrows():
        hover_label = f"Model: {det['ID']}<br>Mục tiêu: {det['Gas']}<br>Cao độ Z: {det['Z']}m"
        fig.add_trace(go.Scatter3d(x=[det['X']], y=[det['Y']], z=[det['Z']], mode='markers+text', 
                                   marker=dict(size=6, color='white'), 
                                   text=[f"{det['ID']}<br>({det['Gas']})"], textposition="top center", textfont=dict(color="white", size=10),
                                   name=det['ID'], hoverinfo="text", hovertext=hover_label))
        
        sx, sy, sz = get_sphere(det['X'], det['Y'], det['Z'], det['Radius'])
        fig.add_trace(go.Surface(x=sx, y=sy, z=sz, opacity=0.15, showscale=False, colorscale=[[0, det['Color']], [1, det['Color']]]))

    for i, (_, obs) in enumerate(df_obs.iterrows()):
        if obs['Type'] == 'Cylinder':
            z_grid, theta = np.mgrid[0:obs['Height']:2j, 0:2*np.pi:20j]
            fig.add_trace(go.Surface(x=obs['Width_Radius']*np.cos(theta)+obs['X'], y=obs['Width_Radius']*np.sin(theta)+obs['Y'], z=z_grid, opacity=1.0, showscale=False, colorscale='Greys', name="Bồn trụ"))
        elif obs['Type'] == 'Box':
            box_2d = obs_polys[i]
            bx, by = box_2d.exterior.xy
            bx, by = list(bx)[:-1], list(by)[:-1] 
            x_box, y_box = bx * 2, by * 2
            z_box = [0]*4 + [obs['Height']]*4
            ii, jj, kk = [7,0,0,0,4,4,6,6,4,0,3,2], [3,4,1,2,5,6,5,2,0,1,6,3], [0,7,2,3,6,7,1,1,5,5,7,6]
            fig.add_trace(go.Mesh3d(x=x_box, y=y_box, z=z_box, i=ii, j=jj, k=kk, color='gray', opacity=1.0, name="Tủ/Kệ"))

    minx, miny, maxx, maxy = room_poly.bounds
    fig.update_layout(
        scene=dict(
            xaxis=dict(range=[minx, maxx], title='X', backgroundcolor="rgb(30,30,30)", gridcolor="gray"),
            yaxis=dict(range=[miny, maxy], title='Y', backgroundcolor="rgb(30,30,30)", gridcolor="gray"),
            zaxis=dict(range=[0, max(rz, 5)], title='Z', backgroundcolor="rgb(30,30,30)", gridcolor="gray"),
            aspectmode='data'
        ),
        paper_bgcolor="rgb(15,15,15)", plot_bgcolor="rgb(15,15,15)",
        margin=dict(l=0, r=0, b=0, t=30), showlegend=False
    )
    return fig

# ==========================================
# 3. WORD REPORT (ĐA LỚP)
# ==========================================
def generate_word_report(figs_dict, img_3d_bytes):
    doc = Document()
    doc.add_heading('BÁO CÁO VÙNG PHỦ KHÍ ĐA LỚP (3D MAPPING)', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Đơn vị thực hiện: CÔNG TY TNHH CÔNG NGHỆ THIẾT BỊ DÒ KHÍ RIKEN VIET').bold = True
    doc.add_paragraph('_' * 60)
    
    doc.add_heading('1. Phân bổ Không gian 3D Tổng thể', level=1)
    if img_3d_bytes:
        doc.add_picture(img_3d_bytes, width=Inches(6.0))
        doc.add_paragraph('Hình 1: Trực quan hóa hệ thống thiết bị và vật cản 3D.').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('2. Phân tích Điểm mù theo Phân hệ Khí', level=1)
    for gas_name, fig in figs_dict.items():
        doc.add_heading(f'Bản đồ giám sát: {gas_name}', level=2)
        img_stream = io.BytesIO()
        fig.savefig(img_stream, format='png', bbox_inches='tight', dpi=150)
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(6.0))
        doc.add_paragraph(f'Hình mô phỏng giao thoa bảo vệ của hệ thống {gas_name}.').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# ==========================================
# 4. TRIGGER KẾT XUẤT
# ==========================================
st.markdown("---")
if st.button("📊 Chạy Mô phỏng Kiến trúc Phức hợp & Tải Báo cáo", use_container_width=True, type='primary'):
    if room_poly is None or edited_dets.empty:
        st.warning("⚠️ Vui lòng nhập đủ tọa độ phòng và danh sách đầu dò!")
    else:
        try:
            obs_polys = create_obstacle_polys(edited_obs)
            collided = check_collision_shapely(edited_dets, obs_polys)
            
            if collided:
                st.error(f"⛔ LỖI VA CHẠM: Đầu dò **{', '.join(collided)}** đang bị đặt nằm bên trong vật cản! Vui lòng chỉnh lại tọa độ X, Y.")
            else:
                with st.spinner('Đang dùng thuật toán Nội suy Raycasting và rà soát bóng mờ đa lớp...'):
                    st.header("3. Phân tích Kết quả Đồ họa")
                    
                    fig_3d = generate_plotly_3d_complex(room_poly, room_z, obs_polys, edited_obs, edited_dets)
                    st.plotly_chart(fig_3d, use_container_width=True)
                    
                    img_3d_bytes = io.BytesIO()
                    try:
                        fig_3d.write_image(img_3d_bytes, format='png', width=800, height=500)
                        img_3d_bytes.seek(0)
                    except: img_3d_bytes = None

                    gas_groups = edited_dets['Gas'].unique()
                    tabs = st.tabs([f"Bản đồ: {g}" for g in gas_groups])
                    generated_figs = {}
                    
                    for i, gas_name in enumerate(gas_groups):
                        with tabs[i]:
                            df_group = edited_dets[edited_dets['Gas'] == gas_name]
                            fig_2d, coverage = generate_2d_plot(room_poly, obs_polys, df_group, gas_name)
                            st.pyplot(fig_2d)
                            if coverage >= 80:
                                st.success(f"✅ Tỷ lệ bao phủ của {gas_name} đạt chuẩn: {coverage:.1f}%")
                            else:
                                st.warning(f"⚠️ Tỷ lệ bao phủ của {gas_name} chỉ đạt: {coverage:.1f}%")
                            generated_figs[gas_name] = fig_2d
                    
                    word_stream = generate_word_report(generated_figs, img_3d_bytes)
                    st.download_button("📄 Tải Báo cáo Tư vấn Chuyên sâu (Word)", word_stream, "Bao_Cao_RikenViet.docx", type="primary")

        except Exception as e:
            st.error(f"Lỗi hệ thống: {e}. Vui lòng kiểm tra lại nét vẽ hoặc thông số nhập liệu.")

# ==========================================
# 5. FOOTER (BẢN QUYỀN TÁC GIẢ)
# ==========================================
st.markdown("""
    <hr style="border: 0; height: 1px; background-image: linear-gradient(to right, rgba(255, 255, 255, 0), rgba(255, 255, 255, 0.2), rgba(255, 255, 255, 0)); margin-top: 50px;">
    <div style="text-align: center; color: #888888; font-size: 14px; padding-bottom: 20px;">
        &copy; 2026 All Rights Reserved.<br>
        Designed and programmed by <b>trggiang</b>.
    </div>
""", unsafe_allow_html=True)
